from concurrent.futures import ThreadPoolExecutor
from flask import Flask, request, jsonify, Response
from llama_index.core import Settings, VectorStoreIndex, Document, StorageContext
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.postprocessor import SimilarityPostprocessor
from transformers import AutoModelForCausalLM, AutoTokenizer, BlipProcessor, BlipForConditionalGeneration
from tqdm import tqdm
from flask_cors import CORS
from PIL import Image
from llama_index.llms.huggingface import HuggingFaceLLM
from docx import Document as DocxDocument

import pandas as pd
import numpy as np
import cv2
import json
import torch
import time
import os
import pickle
import io
import re
from io import BytesIO
import base64
import pytesseract
import PyPDF2
 

os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Initialize Flask app
app = Flask(__name__)
CORS(app)

CACHE_FILE = "chat_cache.json"
CACHE_EXPIRY = 7 * 24 * 60 * 60 # 1 week in seconds

pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'

# Load cache from file
def load_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r") as f:
                item_cache = json.load(f)

            current_time = time.time()
            cache_timestamp = {k: v for k, v in item_cache.items() if current_time - v.get("timestamp", 0) < CACHE_EXPIRY}
            
            return cache_timestamp
        except (json.JSONDecodeError, FileNotFoundError):
            return {}
    return {}

# Save cache to file
def save_cache(cache):
    with open(CACHE_FILE, "w") as f:
        json.dump(cache, f)

# Global cache dictionary
cache = load_cache()

# Initialize Embedding and Vector Index
print("[INFO] Initializing embedding model...")
Settings.embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")
Settings.llm = None
Settings.chunk_size = 1024  
Settings.chunk_overlap = 50

vector_index = None

class CPU_Unpickler(pickle.Unpickler):
    def find_class(self, module, name):
        if module == 'torch.storage' and name == '_load_from_bytes':
            return lambda b: torch.load(io.BytesIO(b), map_location='cpu')
        else:
            return super().find_class(module, name)

def process_entry(entry):
    return Document(
        text=(
            f"Instruction: {entry.get('instruction', '')}\n"
            f"Input: {entry.get('input', '')}\n"
            f"Output: {entry.get('output', '')}"
        )
    )

def save_vector_index(index, file_path):
    with open(file_path, "wb") as f:
        pickle.dump(index, f)
    print(f"[INFO] Vector index saved to {file_path}.")

def load_vector_index(file_path):
    with open(file_path, "rb") as f:
        index = CPU_Unpickler(f).load()  # Use the custom CPU_Unpickler
    print(f"[INFO] Vector index loaded from {file_path}.")
    return index

def read_csv_file(file_path):
    """Read and return the content of a CSV file."""
    try:
        df = pd.read_csv(file_path)
        # Convert the DataFrame to a string representation
        text = df.to_string(index=False)
        return text
    except Exception as e:
        print(f"[ERROR] Failed to read CSV file {file_path}: {str(e)}")
        return None

def read_text_file(file_path):
    """Read and return the content of a text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def read_pdf_file(file_path):
    """Read and return the content of a PDF file."""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text

def read_docx_file(file_path):
    """Read and return the content of a Word document."""
    doc = DocxDocument(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def process_document_file(file_path):
    """Process a document file into a Document object."""
    try:
        if file_path.endswith(".txt"):
            content = read_text_file(file_path)
        elif file_path.endswith(".pdf"):
            content = read_pdf_file(file_path)
        elif file_path.endswith(".docx"):
            content = read_docx_file(file_path)
        elif file_path.endswith(".csv"):
            content = read_csv_file(file_path)
        else:
            print(f"[WARNING] Unsupported file type: {file_path}")
            return None
        return Document(text=content, metadata={"source": file_path})
    except Exception as e:
        print(f"[ERROR] Failed to process file {file_path}: {str(e)}")
        return None

def load_documents(cache_file):
    global vector_index
    print("[INFO] Loading documents into the vector store...")

    try:
        start_time = time.time()

        # Check if the cached file exists
        if os.path.exists(cache_file):
            print("[INFO] Loading vector index from cache...")
            vector_index = load_vector_index(cache_file)
            return

        # Load error and guide JSON files
        with open("instruction-data-errors.json", "r") as errors_file:
            errors_data = json.load(errors_file, strict=False)
        print("[INFO] Loaded 'instruction-data-errors.json'.")

        with open("instruction-data-guide.json", "r") as guide_file:
            guide_data = json.load(guide_file, strict=False)
        print("[INFO] Loaded 'instruction-data-guide.json'.")

        # Combine and process JSON entries
        entries = errors_data + guide_data
        total_entries = len(errors_data) + len(guide_data)
        print(f"[INFO] Total JSON entries to process: {total_entries}")

        with ThreadPoolExecutor(max_workers=30) as executor:
            json_documents = list(
                tqdm(executor.map(process_entry, entries), total=total_entries, desc="[INFO] Processing JSON documents")
            )

        # Process files in the 'documents' directory
        documents_dir = "documents"
        if os.path.exists(documents_dir):
            print(f"[INFO] Reading files from '{documents_dir}' directory...")
            file_paths = [os.path.join(documents_dir, f) for f in os.listdir(documents_dir) if os.path.isfile(os.path.join(documents_dir, f))]
            total_files = len(file_paths)
            print(f"[INFO] Total files to process: {total_files}")

            with ThreadPoolExecutor(max_workers=30) as executor:
                file_documents = list(
                    tqdm(executor.map(process_document_file, file_paths), total=total_files, desc="[INFO] Processing files")
                )
            file_documents = [doc for doc in file_documents if doc is not None]  # Filter out None values
        else:
            print(f"[WARNING] Directory '{documents_dir}' not found. Skipping file processing.")
            file_documents = []

        # Combine JSON and file documents
        all_documents = json_documents + file_documents
        print(f"[INFO] Total documents to index: {len(all_documents)}")

        # Measure time for vector index creation
        vector_start_time = time.time()
        print("[INFO] Creating vector index...")

        storage_context = StorageContext.from_defaults()
        vector_index = VectorStoreIndex.from_documents(all_documents, storage_context=storage_context)

        vector_end_time = time.time()
        print("[INFO] Vector index created successfully.")

        total_time = time.time() - start_time
        vector_time = vector_end_time - vector_start_time

        print(f"[INFO] Total time taken: {total_time:.2f} seconds")
        print(f"[INFO] Time for vector index creation: {vector_time:.2f} seconds")

        save_vector_index(vector_index, cache_file)
        print("[INFO] Vector index saved to cache.")

    except Exception as e:
        print(f"[ERROR] Failed to load documents: {str(e)}")
        raise

load_documents("vector_store_cache.pkl")

# Configure Retriever
print("[INFO] Configuring retriever...")
top_k = 3
retriever = VectorIndexRetriever(index=vector_index, similarity_top_k=top_k)

# Query Engine
print("[INFO] Setting up query engine...")
query_engine = RetrieverQueryEngine(
    retriever=retriever,
    node_postprocessors=[SimilarityPostprocessor(similarity_cutoff=0.8)],
)

print("[INFO] Loading BLIP model...")
processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
image_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")

# Load the model and tokenizer
model_name = "mlx-community/deepseek-r1-distill-qwen-1.5b"

print("[INFO] Loading LLM...")
try:
    # Load model Load the tokenizer
    model = AutoModelForCausalLM.from_pretrained(model_name, torch_dtype=torch.float16, device_map="auto") 
    tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True)
    # Extend Tokenizer Position Embeddings

    print("[INFO] Tokeniser successfully loaded ")
except Exception as e:
    print(f"[ERROR] Failed to load LLM or tokenizer: {str(e)}")
    raise

@app.route('/chat', methods=['POST'])
def chat():
    print("[INFO] Received a new chat request here. Please wait ...")
    try:
        # Parse the incoming JSON request
        data = request.json
	# Log the entire content of the received JSON
        print(f"[INFO] Received data: {data}")
        user_query = data.get("prompt", "")
        image_url = data.get("image_data", None)

        if not user_query:
            print("[WARNING] Query is missing in the request.")
            return jsonify({"error": "Query is required"}), 400

        current_time = time.time()

        # Check if the query is already in the cache
        if user_query in cache and current_time - cache[user_query]["timestamp"] < CACHE_EXPIRY:
            print("[INFO] Returning cached response.")
            return jsonify({"response": cache[user_query]["response"]})

        print(f"[INFO] Query: {user_query}")

        # Retrieve context using RAG
        print("[INFO] Retrieving context using RAG...")
        try:
            response = query_engine.query(user_query)
        except Exception as e:
            print(f"[ERROR] Query execution failed: {str(e)}")
            return jsonify({"error": "Query execution failed. Please try again later."}), 500

        context = "Context:\n"
        num_nodes = min(top_k, len(response.source_nodes))
        for i in range(num_nodes):
            context += response.source_nodes[i].text + "\n\n"

        instruction = (
        "You are ChatFMS, a senior AI expert in CWT Globelink systems. "
        "Provide clear, step-by-step instructions for user queries. "
        "Break down the problem, address potential issues, and recommend the best approach. "
        "Avoid unnecessary details or speculation. Be concise and actionable."
        )        

        prompt = (
        f"[INST]\n{instruction}\n\n"
        f"Context:\n{context}\n\n"
        f"Query: {user_query}\n"
        "1. Break the problem into steps.\n"
        "2. Address potential issues.\n"
        "3. Recommend the best approach.\n"
        "4. Provide concise instructions.\n"
        "[/INST]"
        )        
             
        if image_url and image_url.startswith("data:image"):
            try:
                print("[INFO] Processing base64 image...")
                image_data = base64.b64decode(image_url.split(",")[1])
                image = Image.open(BytesIO(image_data)).convert("RGB")
                image_cv = np.array(image)
                gray = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)
                clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
                gray = clahe.apply(gray)                

                # Apply Adaptive Thresholding
                thresh = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY_INV, 11, 2
                )
                
                # Morph open to remove noise and invert image
                kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,3))
                opening = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel, iterations=1)
                invert = 255 - opening

                print("[INFO] Extracting text from image...")
                extracted_text = pytesseract.image_to_string(invert, config='--psm 6')
                if not extracted_text:
                    extracted_text = ""

                print(f"[INFO] OCR Extracted Text: {extracted_text}")

                print("[INFO] Generating image description...")
                encoding = processor(image, return_tensors="pt")
                outputs = image_model.generate(**encoding)
                caption = processor.decode(outputs[0], skip_special_tokens=True)
                print(f"[INFO] Generated Caption: {caption}")

                response_text = ""

                if extracted_text:
                    response_text = f"The image is {caption}. The text in the image reads: {extracted_text}."
                else:
                    response_text = f"The image is {caption}."

                print(f"[INFO] Combined resultant text: {response_text}")

                print("[INFO] Generating response from LLM...")
                start_time = time.time()  # Start timing
                inputs = tokenizer(prompt, return_tensors="pt")
                outputs = model.generate(
                    input_ids=inputs["input_ids"].to("cpu"),  # Move input to CPU (or GPU if available)
                    max_new_tokens=260
                )
                end_time = time.time()  # End timing
                generated_response = tokenizer.decode(outputs[0], skip_special_tokens=True)

                thinking_keywords = [
                    "I need to figure out", "Let me start by", "I think", "I'm not sure",
                    "maybe", "I wonder", "Wait", "so I'll", "I should", "I guess"
                ]
                if any(keyword.lower() in generated_response.lower() for keyword in thinking_keywords):
                    print("[WARNING] Model generated a thinking process. Falling back to context.")
                    generated_response = context  # Replace with the retrieved context
                clean_response = re.sub(r"\[INST\].*?\[/INST\]\]", "", generated_response, flags=re.DOTALL).strip()
                clean_response = re.sub(r"\[INST\].*?\[/INST\]", "", clean_response, flags=re.DOTALL).strip()
                formatted_response = format_for_apex(response_text + clean_response)
                print(formatted_response)
                # Cache the response
                cache[user_query] = {"response": formatted_response, "timestamp": time.time()}
                save_cache(cache)
                print("[INFO] Response successfully generated.")
                return jsonify({"response": formatted_response})

            except Exception as e:
                print(f"[ERROR] General error: {e}")
                return jsonify({"error": "An unexpected error occurred during image processing"}), 500       

        else:
            print("[INFO] Generating response from LLM...")
            start_time = time.time()  # Start timing
            inputs = tokenizer(prompt, return_tensors="pt")
            outputs = model.generate(
            input_ids=inputs["input_ids"].to("cpu"),  # Move input to CPU (or GPU if available)
            max_new_tokens=260)

            end_time = time.time()  # End timing
            generated_response = tokenizer.decode(outputs[0], skip_special_tokens=True)
            thinking_keywords = [
                "I need to figure out", "Let me start by", "I think", "I'm not sure",
                "maybe", "I wonder", "Wait", "so I'll", "I should", "I guess"
                ]
            if any(keyword.lower() in generated_response.lower() for keyword in thinking_keywords):
                print("[WARNING] Model generated a thinking process. Falling back to context.")
                generated_response = context  # Replace with the retrieved context

            clean_response = re.sub(r"\[INST\].*?\[/INST\]\]", "", generated_response, flags=re.DOTALL).strip()
            clean_response = re.sub(r"\[INST\].*?\[/INST\]", "", clean_response, flags=re.DOTALL).strip()
            formatted_response = format_for_apex(clean_response)
            print(formatted_response)
            # Cache the response
            cache[user_query] = {"response": formatted_response, "timestamp": time.time()}
            save_cache(cache)
            print("[INFO] Response successfully generated.")
            return jsonify({"response": formatted_response})
    
    except Exception as e:
        print(f"[ERROR] An error occurred: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    """
    Endpoint to handle file uploads from Oracle APEX.
    """
    try:
        # Check if a file is included in the request
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']

        # Check if the file has a name
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        # Save the file to the specified directory
        upload_dir = "documents"
        if not os.path.exists(upload_dir):
            os.makedirs(upload_dir)

        file_path = os.path.join(upload_dir, file.filename)
        file.save(file_path)
        print(f"[INFO] File saved to: {file_path}")

        # Reload documents into the vector index
        print("[INFO] Reloading documents...")
        load_documents("vector_store_cache.pkl")

        return jsonify({"message": "File uploaded successfully", "file_path": file_path}), 200

    except Exception as e:
        print(f"[ERROR] Failed to upload file: {str(e)}")
        return jsonify({"error": str(e)}), 500

def format_for_apex(clean_response):
    # Remove content between <think> and </think>, including cases where one tag is missing
    clean_response = re.sub(r'<think>.*?</think>', '', clean_response, flags=re.DOTALL)
    clean_response = re.sub(r'</?think>', '', clean_response)  # Ensure leftover <think> or </think> tags are removed
    
    # Remove leading "ChatFMS:" if present
    clean_response = re.sub(r"^\s*\] ChatFMS:", "", clean_response).strip()

    # Replace markdown bold with HTML strong tags
    clean_response = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", clean_response)

    # Wrap the response in a styled div
    clean_response = f'<div style="font-family: Arial, sans-serif; line-height: 1.6;">{clean_response}</div>'

    # Further wrap the response in another styled div
    clean_response = f"""
        <div style="font-family: Arial, sans-serif; line-height: 1.6; color: #333; padding: 10px; border-radius: 5px; border: 1px solid #ddd;">
            {clean_response}
        </div>
    """
    # Convert numbered lists to HTML ordered lists
    clean_response = clean_response.replace('1. ', '<ol><li>').replace('2. ', '<li>').replace('3. ', '<li>').replace('4. ', '<li>')
    clean_response = clean_response.replace('5. ', '<li>').replace('6. ', '<li>').replace('7. ', '<li>')
    clean_response += '</ol>'

    return clean_response

if __name__ == "__main__":
    print("[INFO] Starting the Flask server on port 5000...")
    app.run(host="0.0.0.0", port=5000)

